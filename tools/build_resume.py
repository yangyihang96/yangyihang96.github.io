from datetime import datetime, timezone
from pathlib import Path
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
PDF_PATH = ASSETS / "Henry_Yang_Biomedical_Engineer_Resume.pdf"
DOCX_PATH = ASSETS / "Henry_Yang_Biomedical_Engineer_Resume.docx"

PROFILE = (
    "Nearly three years of field and workshop service experience at "
    "Nova Biomedical Australia with medical equipment used in hospital and pharmacy settings. "
    "Core strengths include preventive maintenance (PM), fault diagnosis, repair, installation support, "
    "verification, service documentation, and next-use or escalation records."
)

CONTACT_LINE = (
    "Sydney, NSW | 0436 016 660 | yangyihang96@gmail.com | Driver licence | "
    "Available for Sydney field travel | Work rights available for employer verification"
)

SERVICE_BULLETS = [
    "Perform preventive maintenance (PM), fault diagnosis, repair, installation support, verification, and service documentation across field and workshop settings.",
    "Support ventilation, patient monitoring, ultrasound, DEXA, pharmacy automation, X-ray support, and general biomedical equipment.",
    "Document service outcomes with functional checks, performance evidence, or clear escalation status.",
    "Use manufacturer-led checks, service history, and functional evidence to document whether equipment is ready for use, requires follow-up, or needs escalation.",
    "Maintain Simpro work orders, service reports, equipment history, and communication notes so service records remain traceable service evidence.",
]

SCOPE_BULLETS = [
    "Respiratory: V60, V60 Plus, Trilogy; planned service, checks, and troubleshooting preparation.",
    "Monitoring: Avalon, Efficia, HeartStart; service preparation, checks, and handover notes.",
    "Imaging and diagnostics: EPIQ, Affiniti, CX30, CX50, Horizon DEXA, and X-ray support.",
    "Pharmacy automation: BD FIX100, Pyxis, ROWA, workflow support, records, and handover.",
]

SKILL_BULLETS = [
    "PM procedures, fault diagnosis, functional testing, performance verification, service reports, and escalation notes.",
    "Simpro / CMMS, equipment history, equipment identifiers, communication notes, traceability, manufacturer documentation, and Microsoft Office.",
    "Electrical safety testing awareness, test equipment familiarity, service handover, Mandarin Chinese, and professional working proficiency in English.",
]

FIELD_SERVICE_TOOLS = [
    "PM procedures, service manuals, functional testing, performance verification, and clear escalation notes.",
    "Simpro / CMMS, service reports, equipment history, communication notes, handover records, and traceable service evidence.",
    "Electrical safety testing awareness, test equipment familiarity, and manufacturer documentation.",
    "Mandarin Chinese native; English professional working proficiency.",
]

EDUCATION_BULLETS = [
    "Master of Philosophy, The University of Sydney, awarded Jun 2024.",
    "Bachelor of Biomedical Engineering, The University of Sydney, Feb 2017 - Dec 2020.",
    "MPhil thesis: flexible electrodes, impedance measurement, validation evidence, and technical documentation.",
]

OUTCOME_BULLETS = [
    "Supported PM, repair, verification, and documentation across ventilation, monitoring, ultrasound, DEXA, pharmacy automation, and general biomedical equipment in hospital, pharmacy, and workshop settings.",
    "Diagnosed user-reported faults by separating device condition, accessory context, workflow, repair history, reproducible symptoms, procedure-led checks, and post-repair verification.",
    "Supported next-use decisions by documenting whether equipment was ready for use, required follow-up, or needed escalation based on post-service evidence.",
]

ROLE_FIT_BULLETS = [
    "Biomedical Field Service Engineer.",
    "Medical Device Service Engineer.",
    "Clinical Engineering Service Support.",
    "Biomedical Technician / Service Technician.",
]


def wrap_text(text, font_name, font_size, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(c, text, x, y, width, font_name="Helvetica", font_size=8.7, leading=11, color=colors.black):
    c.setFont(font_name, font_size)
    c.setFillColor(color)
    for line in wrap_text(text, font_name, font_size, width):
      c.drawString(x, y, line)
      y -= leading
    return y


def draw_section(c, title, items, x, y, width, title_color):
    c.setFillColor(title_color)
    c.setFont("Helvetica-Bold", 9.6)
    c.drawString(x, y, title.upper())
    y -= 13
    c.setStrokeColor(colors.Color(0.73, 0.80, 0.78))
    c.line(x, y + 5, x + width, y + 5)
    y -= 5
    for item in items:
        y = draw_wrapped(c, f"- {item}", x, y, width, font_size=8.5, leading=10.4, color=colors.Color(0.11, 0.18, 0.17))
        y -= 3
    return y - 6


def build_pdf():
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    width, height = A4
    left = 42
    right = width - 42
    teal = colors.Color(0.06, 0.38, 0.35)
    ink = colors.Color(0.07, 0.13, 0.12)
    muted = colors.Color(0.32, 0.39, 0.38)

    c.setTitle("Henry Yang Biomedical Field Service Engineer Resume")
    c.setAuthor("Yihang Henry Yang")
    c.setSubject("Biomedical field service resume")
    c.setCreator("Yihang Henry Yang")
    c.setProducer("Yihang Henry Yang")
    c.setKeywords("Biomedical Engineer, Field Service Engineer, Medical Device Service, Sydney")

    y = height - 42
    c.setFillColor(teal)
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2, y, "YIHANG (HENRY) YANG")
    y -= 20
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(width / 2, y, "Biomedical Field Service Engineer | Sydney")
    y -= 15
    c.setFillColor(muted)
    c.setFont("Helvetica", 8.5)
    c.drawCentredString(width / 2, y, CONTACT_LINE)
    y -= 18
    c.setStrokeColor(colors.Color(0.73, 0.80, 0.78))
    c.line(left, y, right, y)
    y -= 16

    c.setFillColor(teal)
    c.setFont("Helvetica-Bold", 9.8)
    c.drawString(left, y, "PROFESSIONAL PROFILE")
    y -= 13
    y = draw_wrapped(c, PROFILE, left, y, right - left, font_size=9.1, leading=11.6, color=ink)
    y -= 10

    c.setFillColor(teal)
    c.setFont("Helvetica-Bold", 9.8)
    c.drawString(left, y, "PROFESSIONAL EXPERIENCE")
    y -= 14
    c.setFillColor(ink)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(left, y, "Biomedical Engineer | Nova Biomedical Australia")
    c.setFillColor(muted)
    c.setFont("Helvetica", 8.5)
    c.drawRightString(right, y, "Jul 2023 - Present")
    y -= 12
    c.setFillColor(muted)
    c.setFont("Helvetica", 8.4)
    c.drawString(left, y, "Australia-wide field service / workshop support")
    y -= 12
    for item in SERVICE_BULLETS:
        y = draw_wrapped(c, f"- {item}", left + 8, y, right - left - 8, font_size=8.5, leading=10.3, color=ink)
        y -= 2
    y -= 6

    col_gap = 24
    col_w = (right - left - col_gap) / 2
    y_left = y
    y_right = y
    y_left = draw_section(c, "Equipment and training scope", SCOPE_BULLETS, left, y_left, col_w, teal)
    y_left = draw_section(c, "Technical skills", SKILL_BULLETS, left, y_left, col_w, teal)

    x2 = left + col_w + col_gap
    y_right = draw_section(c, "Earlier healthcare documentation", [
        "Pharmacovigilance Department Assistant, Lundbeck Beijing, Dec 2019 - Feb 2020.",
        "Supported adverse reaction record handling and regulated drug-safety documentation.",
    ], x2, y_right, col_w, teal)
    y_right = draw_section(c, "Education", EDUCATION_BULLETS, x2, y_right, col_w, teal)

    y = min(y_left, y_right) - 4
    c.setStrokeColor(colors.Color(0.73, 0.80, 0.78))
    c.line(left, y + 6, right, y + 6)
    y -= 8
    y = draw_section(c, "Selected service outcomes", OUTCOME_BULLETS, left, y, right - left, teal)

    y_left = y
    y_right = y
    y_left = draw_section(c, "Target roles", ROLE_FIT_BULLETS, left, y_left, col_w, teal)
    y_right = draw_section(c, "Field service tools", FIELD_SERVICE_TOOLS, x2, y_right, col_w, teal)

    footer_y = 34
    c.setStrokeColor(colors.Color(0.73, 0.80, 0.78))
    c.line(left, footer_y + 14, right, footer_y + 14)
    c.setFillColor(muted)
    c.setFont("Helvetica", 7.6)
    c.drawString(left, footer_y, "Updated June 2026.")
    c.showPage()
    c.save()


def set_run_style(run, bold=False, color=None, size=9):
    run.bold = bold
    run.italic = False
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run_style(r, bold=True, color=(17, 54, 90), size=10)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_style(r, size=9)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("YIHANG (HENRY) YANG")
    set_run_style(r, bold=True, color=(17, 54, 90), size=17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Biomedical Field Service Engineer | Sydney")
    set_run_style(r, bold=True, color=(17, 54, 90), size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CONTACT_LINE)
    set_run_style(r, color=(82, 82, 82), size=8.5)

    add_heading(doc, "Professional Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(PROFILE)
    set_run_style(r, size=9.2)

    add_heading(doc, "Professional Experience")
    p = doc.add_paragraph()
    r = p.add_run("Biomedical Engineer | Nova Biomedical Australia")
    set_run_style(r, bold=True, color=(17, 54, 90), size=10)
    r = p.add_run(" | Australia-wide field service / workshop support | Jul 2023 - Present")
    set_run_style(r, color=(82, 82, 82), size=9)
    for item in SERVICE_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "Equipment and Service Scope")
    for item in SCOPE_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "Technical Skills")
    for item in SKILL_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "Earlier Healthcare Documentation")
    add_bullet(doc, "Pharmacovigilance Department Assistant, Lundbeck Beijing, Dec 2019 - Feb 2020.")
    add_bullet(doc, "Supported adverse reaction record handling and regulated drug-safety documentation.")

    add_heading(doc, "Education")
    for item in EDUCATION_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "Selected Service Outcomes")
    for item in OUTCOME_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "Target Roles")
    for item in ROLE_FIT_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "Field Service Tools")
    for item in FIELD_SERVICE_TOOLS:
        add_bullet(doc, item)

    core = doc.core_properties
    core.title = "Henry Yang Biomedical Field Service Engineer Resume"
    core.author = "Yihang Henry Yang"
    core.subject = "Biomedical field service resume"
    core.comments = "Biomedical field service resume for Yihang Henry Yang"
    core.keywords = "Biomedical Engineer, Field Service Engineer, Medical Device Service, Sydney"
    core.last_modified_by = "Yihang Henry Yang"
    created = datetime(2026, 6, 25, tzinfo=timezone.utc)
    core.created = created
    core.modified = created

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    ASSETS.mkdir(exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
